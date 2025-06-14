import os, csv
import requests
from urllib.parse import quote
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import time
import random



proxies = {
  "http": "127.0.0.1:7890", 
  "https": "127.0.0.1:7890"
}

headers = {
    'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
    'accept-language': 'zh-CN,zh;q=0.9',
    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/137.0.0.0 Safari/537.36',
}

cookies = {
    '__Secure-1PSID': 'g.a000xQhGLWSASI9DkCDLAkjKnSSpvuk73vQbpJMPRHn9RBCCsJh0JaLxsX3k8EML8V8MyzhhDAACgYKAf8SARASFQHGX2MiNu-MDXWlor6QCfT2ObdadRoVAUF8yKqgrGIINxgJDX2dyJfXsT0z0076',
}

def readText(path: str):
  # read text
  with open(path, "r") as file:
    data = file.read().strip()
    lines = data.split("\n")
    return lines


def splitName(name: str):
  result = []
  for p1 in name.split(","):
    for p2 in p1.split(":"):
      for p3 in p2.split("["):
        for p4 in p3.split("."):
          result.append(p4)
  result_1 = []
  for p in result:
    if len(p.split(" ")) >= 5:
      result_1.append(p)
  if len(result_1) >= 1:
    return result_1[0].strip()
  return "paper"


def getCitePageUrl(name: str):
  # get cite Page url
  url = "https://scholar.google.com.hk/scholar?hl=zh-CN&as_sdt=0%2C5&q={}&btnG=".format(quote(name))
  response = requests.get(url=url, proxies=proxies, headers=headers)
  soup = BeautifulSoup(response.text, 'lxml')
  flb = soup.find(class_="gs_flb")
  for a in flb.find_all("a"):
    if "被引用" in a.text:
      return "https://scholar.google.com.hk{}".format(a['href'])
  return None


def getCitePageDatas(context, page_index):
  # parse this page
  download_urls = []
  word_data = []
  cite_data = "https://scholar.google.com.hk/scholar?q=info:{}:scholar.google.com/&output=cite&scirp={}&hl=zh-CN"
  soup = BeautifulSoup(context, "lxml")
  mid = soup.find(id="gs_res_ccl_mid")
  divs = mid.find_all("div", recursive=False)
  for i in range(len(divs)):
    # cite word url part
    cid = divs[i]["data-cid"]
    cite_data_url = cite_data.format(cid, (page_index + i))
    response_ = requests.get(url=cite_data_url, headers=headers, proxies=proxies, cookies=cookies)
    soup_ = BeautifulSoup(response_.text, "lxml")
    word_data.append(soup_.find(class_="gs_citr").text)
    # cite pdf url part
    other_page_div = divs[i].find(class_="gs_ri")
    download_url = other_page_div.find(id=cid)
    download_urls.append([download_url.text, download_url["href"]])
    time.sleep(random.randint(1, 3))
  return word_data, download_urls
    

def CiteNextUrl(context):
  # next page
  soup = BeautifulSoup(context, "lxml")
  down = soup.find(id="gs_n")
  td = down.find("td", {"align": "left"})
  a = td.find("a")
  if a is not None:
    return "https://scholar.google.com.hk{}".format(a['href'])
  return None


def getCiteDatas(url: str):
  words = []
  pdfs = []
  response = requests.get(url=url, headers=headers, proxies=proxies, cookies=cookies)
  if "请进行人机身份验证" in response.text:
    print("========= 请重新获取cookies")
    exit()
  page_index = 0
  while True:
    word_data, download_urls = getCitePageDatas(response.text, page_index)
    words += word_data
    pdfs += download_urls
    next_url = CiteNextUrl(response.text)
    if next_url is None:
      break
    response = requests.get(url=next_url, headers=headers, proxies=proxies, cookies=cookies)
    page_index += 1
  return words, pdfs


def saveData(words, pdf_urls, name):
  name = splitName(name)
  file_path = os.path.join(".\\files", name)
  if not os.path.exists(file_path):
    os.makedirs(file_path)
  word_path = os.path.join(file_path, f"{name}.docx")
  csv_path = os.path.join(file_path, f"{name}.cvs")
  with open(csv_path, mode='w', encoding='utf-8', newline='') as file:
    writer = csv.writer(file)
    writer.writerows(pdf_urls)
  
  doc = Document()
  doc.add_paragraph().add_run(f"{name}").bold = True
  doc.add_paragraph("我们的工作被31篇论文引用，如下所示：\n")
  for i in range(len(words)):
    idx = i + 1
    doc.add_paragraph().add_run(f"({idx}){words[i]}").bold = True
    doc.add_paragraph("我们的工作是[]，被引用了  处，如下所示:\n\n\n\n")
  doc.save(word_path)
  

def main():
  for name in readText("./src_paper.txt"):
    print(f"Getting cite: {name}")
    citePageUrl = getCitePageUrl(name=name)
    words, pdfs = getCiteDatas(url=citePageUrl)
    saveData(words=words, pdf_urls=pdfs, name=name)
    print("== Finished.")
    time.sleep(2, 5)
    

if __name__ == "__main__":
  main()